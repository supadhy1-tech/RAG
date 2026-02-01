"""
RAG Document Assistant - Multi-Format Production Backend
Built with FastAPI, ChromaDB, and OpenAI
Supports: PDF, DOCX, TXT, MD, HTML, CSV, XLSX
"""

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import chromadb
from chromadb.config import Settings
from openai import OpenAI
import os
from typing import List, Optional
import PyPDF2
import io
import hashlib
from datetime import datetime
import logging
from dotenv import load_dotenv
from pathlib import Path

# Additional imports for multi-format support
import docx  # python-docx
import csv
import pandas as pd
from bs4 import BeautifulSoup
import markdown

load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI
app = FastAPI(
    title="RAG Document Assistant API - Multi-Format",
    description="Intelligent document Q&A using RAG - Supports PDF, DOCX, TXT, MD, HTML, CSV, XLSX",
    version="2.0.0"
)

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize ChromaDB
chroma_client = chromadb.PersistentClient(path="./chroma_db")

# Create or get collection
try:
    collection = chroma_client.get_collection(name="documents")
except:
    collection = chroma_client.create_collection(
        name="documents",
        metadata={"hnsw:space": "cosine"}
    )

# OpenAI Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
client = OpenAI(api_key=OPENAI_API_KEY)

# Supported file types
SUPPORTED_EXTENSIONS = {
    '.pdf': 'PDF Document',
    '.docx': 'Word Document',
    '.txt': 'Text File',
    '.md': 'Markdown File',
    '.html': 'HTML File',
    '.htm': 'HTML File',
    '.csv': 'CSV File',
    '.xlsx': 'Excel Spreadsheet',
    '.xls': 'Excel Spreadsheet',
}

# Pydantic models
class QueryRequest(BaseModel):
    question: str
    top_k: int = 3

class QueryResponse(BaseModel):
    answer: str
    sources: List[dict]
    confidence: float
    latency_ms: float

class DocumentInfo(BaseModel):
    id: str
    filename: str
    file_type: str
    upload_time: str
    chunk_count: int

# ==================== TEXT EXTRACTION FUNCTIONS ====================

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF bytes"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX bytes"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT bytes"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with any common encoding")
    except Exception as e:
        logger.error(f"Error extracting TXT text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from TXT: {str(e)}")

def extract_text_from_markdown(file_content: bytes) -> str:
    """Extract text from Markdown bytes"""
    try:
        text = file_content.decode('utf-8')
        # Convert markdown to HTML, then extract text
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text(separator='\n')
    except Exception as e:
        logger.error(f"Error extracting Markdown text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from Markdown: {str(e)}")

def extract_text_from_html(file_content: bytes) -> str:
    """Extract text from HTML bytes"""
    try:
        html = file_content.decode('utf-8')
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n')
        # Clean up excessive whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        logger.error(f"Error extracting HTML text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from HTML: {str(e)}")

def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV bytes"""
    try:
        text_content = file_content.decode('utf-8')
        csv_reader = csv.reader(io.StringIO(text_content))
        
        rows = []
        for row in csv_reader:
            rows.append(" | ".join(row))
        
        return "\n".join(rows)
    except Exception as e:
        logger.error(f"Error extracting CSV text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from CSV: {str(e)}")

def extract_text_from_excel(file_content: bytes) -> str:
    """Extract text from Excel bytes (XLSX/XLS)"""
    try:
        df_dict = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
        
        text = ""
        for sheet_name, df in df_dict.items():
            text += f"\n=== Sheet: {sheet_name} ===\n"
            text += df.to_string(index=False)
            text += "\n"
        
        return text
    except Exception as e:
        logger.error(f"Error extracting Excel text: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from Excel: {str(e)}")

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Route to appropriate extraction function based on file extension"""
    ext = Path(filename).suffix.lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.md': extract_text_from_markdown,
        '.html': extract_text_from_html,
        '.htm': extract_text_from_html,
        '.csv': extract_text_from_csv,
        '.xlsx': extract_text_from_excel,
        '.xls': extract_text_from_excel,
    }
    
    extractor = extractors.get(ext)
    if not extractor:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Supported types: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )
    
    return extractor(file_content)

# ==================== UTILITY FUNCTIONS ====================

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size
        chunk = text[start:end]

        # Try to break at sentence boundary
        if end < text_length:
            last_period = chunk.rfind('.')
            if last_period > chunk_size * 0.5:
                chunk = chunk[:last_period + 1]
                end = start + last_period + 1

        chunks.append(chunk.strip())
        start = end - overlap

    return [c for c in chunks if len(c.strip()) > 50]

def generate_doc_id(filename: str, content: bytes) -> str:
    """Generate unique document ID"""
    hash_content = hashlib.md5(content).hexdigest()
    return f"{filename}_{hash_content[:8]}"

def get_file_type(filename: str) -> str:
    """Get file type description from filename"""
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext, "Unknown")

# ==================== API ENDPOINTS ====================

@app.get("/")
async def root():
    return {
        "message": "RAG Document Assistant API - Multi-Format",
        "version": "2.0.0",
        "status": "operational",
        "supported_formats": list(SUPPORTED_EXTENSIONS.keys())
    }

@app.get("/supported-formats")
async def supported_formats():
    """List all supported file formats"""
    return {
        "formats": [
            {"extension": ext, "description": desc}
            for ext, desc in SUPPORTED_EXTENSIONS.items()
        ]
    }

@app.post("/upload", response_model=DocumentInfo)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document (supports multiple formats)"""
    start_time = datetime.now()

    # Validate file extension
    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Supported types: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    try:
        content = await file.read()
        doc_id = generate_doc_id(file.filename, content)
        
        logger.info(f"Extracting text from {file.filename} ({SUPPORTED_EXTENSIONS[ext]})")
        text = extract_text_from_file(content, file.filename)

        if not text.strip():
            raise HTTPException(status_code=400, detail="No text found in file")

        logger.info(f"Chunking text from {file.filename}")
        chunks = chunk_text(text)

        logger.info(f"Storing {len(chunks)} chunks in database")
        collection.add(
            documents=chunks,
            ids=[f"{doc_id}_chunk_{i}" for i in range(len(chunks))],
            metadatas=[{
                "filename": file.filename,
                "file_type": SUPPORTED_EXTENSIONS[ext],
                "doc_id": doc_id,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "upload_time": datetime.now().isoformat()
            } for i in range(len(chunks))]
        )

        elapsed = (datetime.now() - start_time).total_seconds() * 1000
        logger.info(f"Document processed in {elapsed:.2f}ms")

        return DocumentInfo(
            id=doc_id,
            filename=file.filename,
            file_type=SUPPORTED_EXTENSIONS[ext],
            upload_time=datetime.now().isoformat(),
            chunk_count=len(chunks),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")

@app.post("/query", response_model=QueryResponse)
async def query_documents(request: QueryRequest):
    """Query documents using RAG"""
    start_time = datetime.now()

    if not OPENAI_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable.",
        )

    try:
        logger.info(f"Searching for: {request.question}")
        results = collection.query(query_texts=[request.question], n_results=request.top_k)

        if not results['documents'][0]:
            return QueryResponse(
                answer="I couldn't find any relevant information in the uploaded documents.",
                sources=[],
                confidence=0.0,
                latency_ms=(datetime.now() - start_time).total_seconds() * 1000,
            )

        context_chunks = results['documents'][0]
        sources = []
        for i, (chunk, metadata, distance) in enumerate(
            zip(context_chunks, results['metadatas'][0], results['distances'][0])
        ):
            sources.append({
                "chunk": chunk[:200] + "..." if len(chunk) > 200 else chunk,
                "filename": metadata.get("filename", "Unknown"),
                "file_type": metadata.get("file_type", "Unknown"),
                "relevance": float(1 - distance),
            })

        context = "\n\n".join([f"[Source {i+1} - {sources[i]['filename']}]: {chunk}" for i, chunk in enumerate(context_chunks)])

        logger.info("Generating answer with GPT")
        system_prompt = """You are a helpful AI assistant that answers questions based on provided document excerpts.

Rules:
- Only use information from the provided sources
- If the sources don't contain enough information, say so clearly
- Cite which source(s) you're using (e.g., "According to Source 1...")
- Be concise but thorough
- If asked about something not in the sources, acknowledge the limitation"""

        user_prompt = f"""Context from documents:

{context}

Question: {request.question}

Please provide a clear, accurate answer based on the context above."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=500,
        )

        answer = response.choices[0].message.content

        avg_relevance = sum(s["relevance"] for s in sources) / len(sources)
        confidence = min(avg_relevance * 1.2, 1.0)

        elapsed = (datetime.now() - start_time).total_seconds() * 1000
        logger.info(f"Query completed in {elapsed:.2f}ms")

        return QueryResponse(
            answer=answer,
            sources=sources,
            confidence=confidence,
            latency_ms=elapsed,
        )

    except Exception as e:
        logger.error(f"Error querying documents: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to query documents: {str(e)}")

@app.get("/documents", response_model=List[DocumentInfo])
async def list_documents():
    """List all uploaded documents"""
    try:
        all_data = collection.get()
        docs = {}
        for metadata in all_data['metadatas']:
            doc_id = metadata.get("doc_id")
            if doc_id not in docs:
                docs[doc_id] = {
                    "id": doc_id,
                    "filename": metadata.get("filename"),
                    "file_type": metadata.get("file_type", "Unknown"),
                    "upload_time": metadata.get("upload_time"),
                    "chunk_count": 0,
                }
            docs[doc_id]["chunk_count"] += 1
        return list(docs.values())
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to list documents")

@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document and all its chunks"""
    try:
        results = collection.get(where={"doc_id": doc_id})

        if not results['ids']:
            raise HTTPException(status_code=404, detail="Document not found")

        collection.delete(ids=results['ids'])
        return {"message": f"Deleted document {doc_id} and {len(results['ids'])} chunks"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete document")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "openai_configured": bool(OPENAI_API_KEY),
        "documents_count": len(collection.get()['ids']),
        "supported_formats": list(SUPPORTED_EXTENSIONS.keys())
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
